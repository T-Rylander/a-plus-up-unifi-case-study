#!/usr/bin/env python3
"""Extract all content from Word and Excel documents."""

import os
from pathlib import Path
from docx import Document
from openpyxl import load_workbook

def extract_docx(filepath):
    """Extract all text, tables, and formatting from a DOCX file."""
    content = []
    doc = Document(filepath)
    
    # Extract paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            content.append(para.text)
    
    # Extract tables
    for table in doc.tables:
        table_content = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            table_content.append(" | ".join(row_data))
        content.extend(table_content)
    
    return "\n".join(content)

def extract_xlsx(filepath):
    """Extract all content from an XLSX file."""
    content = []
    wb = load_workbook(filepath)
    
    for sheet_name in wb.sheetnames:
        content.append(f"\n=== Sheet: {sheet_name} ===\n")
        ws = wb[sheet_name]
        
        for row in ws.iter_rows(values_only=True):
            # Filter out None values and convert to strings
            row_data = [str(cell) if cell is not None else "" for cell in row]
            content.append(" | ".join(row_data))
    
    return "\n".join(content)

def main():
    attachments_path = Path("d:/Repos/a-plus-up-unifi-case-study/docs/attachments")
    
    files = [
        "Additional considerations addendum AUP.docx",
        "AUP Network Overhaul Blueprint.docx",
        "Comprehensive WiFi Assessment Summary V2.1.docx",
        "Project Dossier AplusUP.docx",
        "UpFront Cost est - AplusUP Charter.xlsx",
        "UPS Monitoring Guide.docx",
        "Verkada Camera Migration.docx",
        "VoIP migration and Configuration.docx"
    ]
    
    output = []
    
    for filename in files:
        filepath = attachments_path / filename
        
        output.append("=" * 100)
        output.append(f"\nFILE: {filename}\n")
        output.append("=" * 100)
        output.append("")
        
        try:
            if filename.endswith('.docx'):
                content = extract_docx(str(filepath))
            elif filename.endswith('.xlsx'):
                content = extract_xlsx(str(filepath))
            else:
                content = "[Unknown file format]"
            
            output.append(content)
        except Exception as e:
            output.append(f"[ERROR extracting {filename}: {str(e)}]")
        
        output.append("\n")
    
    # Print all output
    full_output = "\n".join(output)
    print(full_output)
    
    # Also save to a file
    output_file = attachments_path / "EXTRACTED_CONTENT.txt"
    with open(output_file, 'w', encoding='utf-8') as f:
        f.write(full_output)
    
    print(f"\n\n[Output also saved to: {output_file}]")

if __name__ == "__main__":
    main()
